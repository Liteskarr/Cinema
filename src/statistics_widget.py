from collections import defaultdict
from datetime import datetime, timedelta

from PyQt5.QtCore import (Qt)
from PyQt5.QtGui import (QFont)
from PyQt5.QtWidgets import (QWidget,
                             QGridLayout,
                             QLabel,
                             QDateTimeEdit,
                             QSizePolicy,
                             QPushButton,
                             QFileDialog)
import xlsxwriter
from docx import Document

from src.ipage import IPage
from src.context_locator import ContextLocator


class StatisticsWidget(QWidget, IPage):
    def __init__(self):
        super().__init__()
        self._configure_ui()

    def _configure_ui(self):
        self.setFont(QFont('Sans', 12))
        self.grid = QGridLayout(self)

        self.from_datetime_edit = QDateTimeEdit(self)
        self.from_datetime_edit.setDateTime(datetime.now() - timedelta(days=30))
        self.grid.addWidget(self.from_datetime_edit, 0, 0)

        self.split_label = QLabel(self)
        self.split_label.setText(' ---> ')
        self.split_label.setAlignment(Qt.AlignCenter)
        self.split_label.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Minimum)
        self.grid.addWidget(self.split_label, 0, 1)

        self.to_datetime_edit = QDateTimeEdit(self)
        self.to_datetime_edit.setDateTime(datetime.now())
        self.grid.addWidget(self.to_datetime_edit, 0, 2)

        self.timetable_button = QPushButton(self)
        self.timetable_button.clicked.connect(self._handle_timetable_buttons_click)
        self.timetable_button.setText('Расписание сеансов за выбранный период (*.docx)')
        self.timetable_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.grid.addWidget(self.timetable_button, 1, 0, 1, 3)

        self.income_button = QPushButton(self)
        self.income_button.clicked.connect(self._handle_income_button_click)
        self.income_button.setText('Табличные данные за выбранный период (*.xlsx)')
        self.income_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.grid.addWidget(self.income_button, 2, 0, 1, 3)

    def _handle_timetable_buttons_click(self):
        document = Document()
        show_format = ContextLocator.get_context().datetime_show_format
        from_date = self.from_datetime_edit.dateTime().toPyDateTime()
        to_date = self.to_datetime_edit.dateTime().toPyDateTime()
        document.add_heading(f'Расписание всех сеансов с {from_date.strftime(show_format)} '
                             f'до {to_date.strftime(show_format)}.', 0)
        request = """
        SELECT films.name, cinemas.name, halls.name, sessions.datetime, films.duration, sessions.cost 
        FROM sessions
        INNER JOIN films ON sessions.film = films.id
        INNER JOIN halls ON sessions.hall = halls.id
        INNER JOIN cinemas ON halls.cinema = cinemas.id
        WHERE
            sessions.datetime BETWEEN ? AND ?
        ORDER BY sessions.datetime ASC;
        """
        connection = ContextLocator.get_context().connection
        cursor = connection.cursor()
        for film, cinema, hall, date, duration, cost in cursor.execute(request,
                                                                       (from_date, to_date)).fetchall():
            date = datetime.strptime(date[:-3], ContextLocator.get_context().datetime_packing_format)
            duration = timedelta(minutes=duration)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f'Дата начала: {date.strftime(show_format)}\n')
            run.bold = True
            run = paragraph.add_run(f'Продолжительность: {duration.seconds // 3600:02d}:'
                                    f'{duration.seconds // 60 % 60:02d}\n')
            run.bold = True
            run = paragraph.add_run(f'Дата конца: {(date + duration).strftime(show_format)}\n')
            run.bold = True
            run = paragraph.add_run(f'Кинотеатр: {cinema}\n')
            run = paragraph.add_run(f'Зал: {hall}\n')
            run = paragraph.add_run(f'Фильм: {film}\n')
        filename, ok = QFileDialog.getSaveFileName(self, 'Выберите файл', '', 'Документ (*.docx);;Все (*.*)')
        if ok:
            document.save(filename)

    def _get_films_income(self, from_date: datetime, to_date: datetime) -> list:
        connection = ContextLocator.get_context().connection
        cursor = connection.cursor()
        films = dict([(p[0], p) for p in cursor.execute("SELECT * FROM films;").fetchall()])
        films_income = defaultdict(int)
        request = """
        SELECT films.id, sessions.cost FROM tickets
        INNER JOIN sessions ON tickets.session = sessions.id
        INNER JOIN films ON sessions.film = films.id
        WHERE
            sessions.datetime BETWEEN ? AND ?;
        """
        for film, cost in cursor.execute(request, (from_date, to_date)).fetchall():
            films_income[film] += cost
        return [(f'{films[f][1]} ({films[f][2] // 60 :02d}:{films[f][2] % 60:02d})', films_income[f])for f in films]

    def _get_cinemas_income(self, from_date: datetime, to_date: datetime) -> list:
        connection = ContextLocator.get_context().connection
        cursor = connection.cursor()
        cinemas = dict([(p[0], p) for p in cursor.execute("SELECT * FROM cinemas;").fetchall()])
        cinemas_income = defaultdict(int)
        request = """
        SELECT cinemas.id, sessions.cost FROM tickets
        INNER JOIN sessions ON sessions.id = tickets.session
        INNER JOIN halls ON halls.id = sessions.hall
        INNER JOIN cinemas ON cinemas.id = halls.cinema
        WHERE
            sessions.datetime BETWEEN ? AND ?;
        """
        for cinema, cost in cursor.execute(request, (from_date, to_date)).fetchall():
            cinemas_income[cinema] += cost
        return [(cinemas[c][1], cinemas_income[c]) for c in cinemas]

    def _handle_income_button_click(self):
        filename, ok = QFileDialog.getSaveFileName(self, 'Выбрать файл.', '', 'Документ (*.xlsx);;Все (*.*)')
        if not ok:
            return
        from_date = self.from_datetime_edit.dateTime().toPyDateTime()
        to_date = self.to_datetime_edit.dateTime().toPyDateTime()
        workbook = xlsxwriter.Workbook(filename)
        worksheet = workbook.add_worksheet('Кинотеатры')
        cinemas_income = self._get_cinemas_income(from_date, to_date)
        worksheet.write(0, 0, 'Кинотеатр')
        worksheet.write_column('A2', map(lambda x: x[0], cinemas_income))
        worksheet.write(0, 1, 'Доход')
        worksheet.write_column('B2', map(lambda x: x[1], cinemas_income))
        worksheet = workbook.add_worksheet('Фильмы')
        films_income = self._get_films_income(from_date, to_date)
        worksheet.write(0, 0, 'Фильм')
        worksheet.write_column('A2', map(lambda x: x[0], films_income))
        worksheet.write(0, 1, 'Доход')
        worksheet.write_column('B2', map(lambda x: x[1], films_income))
        workbook.close()

    def get_new_page(self):
        raise NotImplementedError()
